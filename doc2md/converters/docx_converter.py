from __future__ import annotations

from doc2md.converters import register
from doc2md.models.document import Block, BlockType, Document
from doc2md.utils.image_handler import save_image


@register(".docx")
def convert_docx(path: str) -> Document:
    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx is required for .docx files: pip install python-docx")

    docx = DocxDocument(path)
    doc = Document(metadata={"source": path})

    # metadata
    try:
        core_props = docx.core_properties
        doc.metadata["title"] = core_props.title or ""
        doc.metadata["author"] = core_props.author or ""
    except Exception:
        pass

    para_index = 0
    table_index = 0
    body = docx.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_index < len(docx.paragraphs):
                para = docx.paragraphs[para_index]
                para_index += 1
                block = _convert_paragraph(para)
                if block:
                    doc.blocks.append(block)
            else:
                para_index += 1

        elif tag == "tbl":
            if table_index < len(docx.tables):
                table = docx.tables[table_index]
                table_index += 1
                # Skip paragraphs that belong to this table
                while para_index < len(docx.paragraphs) and _paragraph_index_in_table(docx.paragraphs[para_index], table):
                    para_index += 1
                md_table = _convert_table(table)
                if md_table:
                    doc.blocks.append(Block(type=BlockType.TABLE, content=md_table))

    return doc


def _convert_paragraph(para) -> Block | None:
    text = para.text.strip()
    if not text and not para.runs:
        return None

    style_name = para.style.name.lower() if para.style and para.style.name else ""
    text = para.text.strip()

    if not text:
        return None

    # heading style (includes "Title", "Heading 1-6", "标题")
    if "title" in style_name or "heading" in style_name or "标题" in style_name:
        try:
            level = int(style_name.replace("title", "").replace("heading ", "").replace("标题 ", "").strip())
        except ValueError:
            level = 1
        return Block(type=BlockType.HEADING, level=min(max(level, 1), 6), content=text)

    # list style
    if "list" in style_name:
        level = 1
        if "list " in style_name or "列表 " in style_name:
            try:
                level = int(style_name.replace("list ", "").replace("列表 ", ""))
            except ValueError:
                level = 1
        return Block(type=BlockType.LIST_ITEM, level=level, content=text)

    # code style
    if style_name in ("code", "Code", "代码"):
        return Block(type=BlockType.CODE_BLOCK, content=text, metadata={"language": ""})

    # detect heading from font size (for docs not using heading styles)
    if para.runs and para.runs[0].font.size:
        font_size = para.runs[0].font.size.pt
        if font_size >= 16:
            return Block(type=BlockType.HEADING, level=1 if font_size >= 22 else 2, content=text)

    # preserve inline formatting (bold/italic)
    formatted = _preserve_formatting(para)
    return Block(type=BlockType.PARAGRAPH, content=formatted)


def _is_all_bold(para) -> bool:
    if not para.runs:
        return False
    return all(r.bold for r in para.runs if r.text.strip())


def _convert_table(table) -> str:
    rows_data = []
    col_count = 0

    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        col_count = max(col_count, len(cells))
        rows_data.append(cells)

    if not rows_data:
        return ""

    # pad all rows to col_count
    for row in rows_data:
        row.extend([""] * (col_count - len(row)))

    lines = ["|" + "|".join(rows_data[0]) + "|"]
    lines.append("|" + "|".join(["---"] * col_count) + "|")
    for row in rows_data[1:]:
        lines.append("|" + "|".join(row) + "|")

    return "\n".join(lines)


def _paragraph_index_in_table(para, table) -> bool:
    """Check if a paragraph belongs to a table."""
    try:
        tbl_element = table._tbl
        parent = para._element.getparent()
        while parent is not None:
            if parent is tbl_element:
                return True
            parent = parent.getparent()
    except Exception:
        pass
    return False


def _preserve_formatting(para) -> str:
    """Convert runs to Markdown with basic inline formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        # Escape special Markdown characters
        text = text.replace("\\", "\\\\").replace("`", "\\`")

        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"

        parts.append(text)

    if not parts:
        return para.text
    return " ".join(parts)
