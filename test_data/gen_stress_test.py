"""Generate stress test document for doc2md."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Title ──
title = doc.add_heading("doc2md 边界压力测试报告", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("版本：2.0.0 | 日期：2025-12-20 | 分类：内部测试")

# ── 1. Deep headings ──
doc.add_heading("第一章 多层次标题结构", level=1)
doc.add_paragraph("本文档用于测试 doc2md 转换器的极限能力。")
doc.add_heading("1.1 二级标题", level=2)
doc.add_heading("1.1.1 三级标题", level=3)
doc.add_heading("1.1.1.1 四级标题", level=4)
doc.add_heading("1.1.1.1.1 五级标题", level=5)
doc.add_heading("1.1.1.1.1.1 六级标题", level=6)
doc.add_paragraph("六级标题之后的内容。")

# ── 2. Complex tables ──
doc.add_heading("第二章 复杂表格测试", level=1)
doc.add_heading("2.1 标准数据表", level=2)
t = doc.add_table(rows=6, cols=5)
t.style = "Table Grid"
for i, h in enumerate(["产品名称", "Q1(万)", "Q2(万)", "Q3(万)", "Q4(万)"]):
    t.cell(0, i).text = h
data = [
    ["智能客服", "1,280", "1,450", "1,670", "2,010"],
    ["金融风控", "980", "1,120", "1,350", "1,580"],
    ["数据中台", "2,100", "2,350", "2,680", "3,120"],
    ["AI 训练平台", "750", "890", "1,020", "1,340"],
    ["合计", "5,110", "5,810", "6,720", "8,050"],
]
for r, row in enumerate(data):
    for c, val in enumerate(row):
        t.cell(r+1, c).text = val

doc.add_paragraph("")
doc.add_heading("2.2 大表格压力测试 (20×30)", level=2)
big = doc.add_table(rows=30, cols=20)
big.style = "Table Grid"
for r in range(30):
    for c in range(20):
        big.cell(r, c).text = f"H{c}" if r == 0 else f"D{r:02d}C{c:02d}"

# ── 3. Lists ──
doc.add_heading("第三章 列表结构", level=1)
doc.add_heading("3.1 混合列表", level=2)
for item in ["需求分析", "系统设计", "编码实现", "测试部署"]:
    doc.add_paragraph(item, style="List Number")
for tool in ["Python 3.12", "FastAPI", "PostgreSQL", "Docker"]:
    doc.add_paragraph(tool, style="List Bullet")

# ── 4. Formatting ──
doc.add_heading("第四章 文本格式测试", level=1)
p = doc.add_paragraph()
r = p.add_run("粗体文本 ")
r.bold = True
r = p.add_run("斜体文本 ")
r.italic = True
r = p.add_run("下划线文本 ")
r.underline = True
r = p.add_run("删除线文本")
r.font.strike = True

p2 = doc.add_paragraph()
r = p2.add_run("红色字体")
r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
r = p2.add_run(" 蓝色字体")
r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

p3 = doc.add_paragraph()
r = p3.add_run("H")
r = p3.add_run("2")
r.font.subscript = True
r = p3.add_run("O 上标：m")
r = p3.add_run("2")
r.font.superscript = True

# ── 5. Long paragraph ──
doc.add_heading("第五章 长段落压力", level=1)
doc.add_paragraph("压力测试：" * 300 + "长段落结束。")

# ── 6. Mixed content ──
doc.add_heading("第六章 综合内容", level=1)
p = doc.add_paragraph()
r = p.add_run("【重要】")
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
p.add_run(" 根据规划要求，所有系统必须达到 ")
r = p.add_run("99.99%")
r.bold = True
p.add_run(" 的可用性标准。")

doc.add_paragraph("第一段：微服务架构设计，每个服务独立部署、独立扩展。")
doc.add_paragraph("第二段：采用 Kubernetes + Istio 实现流量管理和安全策略。")
doc.add_paragraph("第三段：数据存储采用混合方案，热数据使用 Redis 缓存。")

# ── 7. Wide table ──
doc.add_heading("第七章 超宽表格测试", level=1)
wide = doc.add_table(rows=4, cols=15)
wide.style = "Table Grid"
for r in range(4):
    for c in range(15):
        wide.cell(r, c).text = f"参{c+1}" if r == 0 else f"值{r}_{c+1}"

# ── 8. Special chars ──
doc.add_heading("第八章 特殊字符", level=1)
doc.add_paragraph("特殊字符：© ® ™ ∞ ≈ ≠ ≤ ≥ ± × ÷ ← ↑ → ↓ ★ ☆ ♠ ♣ ♥ ♦")
doc.add_paragraph("Emoji：\U0001f600 \U0001f680 \U0001f31f \U0001f525 ✅ ❌ ⚠️ \U0001f4a1 \U0001f4ca \U0001f3af")
doc.add_paragraph("代码片段：print(\"hello\") 和 const x = await fetch(\"/api\")")
doc.add_heading("8.1 空段落", level=2)
doc.add_paragraph("")
doc.add_paragraph("空段落之后的正文。")

# ── Save ──
path = r"G:\my code\claude code\doc2md\test_data\stress_test.docx"
doc.save(path)
size = os.path.getsize(path)
print(f"Created: {path}")
print(f"Size: {size} bytes ({size/1024:.1f} KB)")

d = Document(path)
total = sum(len(p.text) for p in d.paragraphs)
print(f"Total chars: {total}")
print(f"Tables: {len(d.tables)}")
